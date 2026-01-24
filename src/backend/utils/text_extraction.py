import mimetypes
import fitz
from PIL import Image
from docx import Document as DocxDocument
from pypdf import PdfReader
import io

def extract_text_from_file(file_bytes, mime_type):
    if mime_type == 'application/pdf':
        # Try text extraction first
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            text = "\n".join([page.extract_text() for page in reader.pages])
            if text.strip():
                return text
        except Exception:
            pass
        # Else OCR
        import easyocr
        ocr_reader = easyocr.Reader(['en', 'mk'], gpu=False)  # English and Macedonian, CPU
        doc = fitz.Document(stream=file_bytes, filetype="pdf")
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap()
            img_bytes = pix.tobytes()
            img = Image.open(io.BytesIO(img_bytes))
            results = ocr_reader.readtext(img)
            text += " ".join([res[1] for res in results]) + "\n"
        return text
    elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
        doc = DocxDocument(io.BytesIO(file_bytes))
        return "\n".join([para.text for para in doc.paragraphs])
    elif mime_type.startswith('text/'):
        return file_bytes.decode("utf-8")
    else:
        raise ValueError("Unsupported file type")